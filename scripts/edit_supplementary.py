"""
Polish SupplementaryData_AgeAgeing_v0418.docx for Age and Ageing submission.
Saves edited copy; use Word Review > Compare Documents for tracked changes.
"""

import shutil
import docx

SRC = r"C:\Users\xc77\Dropbox\Grip Strength and Cognitive Trajectories\SupplementaryData_AgeAgeing_v0418.docx"
DST = r"C:\Users\xc77\Dropbox\Grip Strength and Cognitive Trajectories\SupplementaryData_AgeAgeing_v0418_edited.docx"

shutil.copy(SRC, DST)
doc = docx.Document(DST)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def full_text(para):
    return "".join(r.text for r in para.runs)


def replace_in_para(para, old, new):
    ft = full_text(para)
    if old not in ft:
        return False
    new_ft = ft.replace(old, new, 1)
    if para.runs:
        para.runs[0].text = new_ft
        for r in para.runs[1:]:
            r.text = ""
    return True


def replace_in_cell(cell, old, new):
    """Replace text inside a table cell (across its paragraphs)."""
    for para in cell.paragraphs:
        if replace_in_para(para, old, new):
            return True
    return False


log = []

def edit(para, old, new, label):
    if replace_in_para(para, old, new):
        log.append(("DONE", label))
        return True
    return False


# ---------------------------------------------------------------------------
# Text paragraph edits
# ---------------------------------------------------------------------------

all_paras = list(doc.paragraphs)

for para in all_paras:

    # 1. "one to six groups" -> "one- to six-group models" (match main text edit)
    edit(para,
         "we tested models specifying one to six groups",
         "we tested one- to six-group models",
         "Methods: 'one to six groups' -> 'one- to six-group models' (consistent with main text)")

    # 2. Redundant double use of "criterion"
    edit(para,
         "The Bayesian Information Criterion (BIC) was used as the primary model selection criterion.",
         "The Bayesian Information Criterion (BIC) guided model selection.",
         "Methods: removed redundant double 'criterion'")

    # 3. Remove verbose redundancy about posterior probabilities
    edit(para,
         ("The estimated group proportions aligned closely with the observed proportions of "
          "individuals assigned to each group based on posterior probabilities, indicating "
          "good classification consistency."),
         ("The estimated group proportions aligned closely with the observed assignment "
          "proportions, indicating good classification consistency."),
         "Methods: simplified redundant posterior-probability clause")

    # 4. Kendall's W -- consistent with main text terminology
    edit(para,
         "Kendall's coefficient of concordance",
         "Kendall's W (coefficient of concordance)",
         "Methods: 'Kendall's coefficient of concordance' -> 'Kendall's W' (consistent with main text)")

    # 5. Table S2 note: imprecise 'chronic conditions' -> matches Fig S2 note wording
    edit(para,
         ", chronic conditions, and global cognition.",
         ", 13 physician-diagnosed chronic conditions, and baseline global cognition.",
         "Table S2 note: added '13 physician-diagnosed' and 'baseline' for precision (consistent with Figure S2 note)")

    # 6. Figure S2 note: American 'standardized' -> British 'standardised'
    edit(para,
         "standardized regression coefficients",
         "standardised regression coefficients",
         "Figure S2 note: British English 'standardized' -> 'standardised'")


# ---------------------------------------------------------------------------
# Table S1 cell edits
# ---------------------------------------------------------------------------

t1 = doc.tables[0]

# Fix 7: Currently smoke, Men Diff: "0.017" should be "1.7%"
# Row 20, col 7
cell_smoke_diff = t1.rows[20].cells[7]
if replace_in_cell(cell_smoke_diff, "0.017", "1.7%"):
    log.append(("DONE",
                "Table S1: Men 'Currently smoke' Diff corrected from '0.017' to '1.7%' "
                "(decimal proportion -> percentage, consistent with all other Diff. values)"))
else:
    log.append(("MISS", "Table S1: Could not find '0.017' in smoke Diff cell"))


# ---------------------------------------------------------------------------
# Table S2 cell edits
# ---------------------------------------------------------------------------

t2 = doc.tables[1]

# Fix 8: Observations count "1 ,653" -> "1,653" (spurious space before comma)
fixed_obs = False
for row in t2.rows:
    for cell in row.cells:
        if replace_in_cell(cell, "1 ,653", "1,653"):
            fixed_obs = True
if fixed_obs:
    log.append(("DONE", "Table S2: Fixed spurious space in observation count '1 ,653' -> '1,653'"))
else:
    log.append(("MISS", "Table S2: Could not find '1 ,653'"))


# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------

doc.save(DST)


# ---------------------------------------------------------------------------
# Report
# ---------------------------------------------------------------------------

SEP = "-" * 65
print(SEP)
print("SUPPLEMENTARY EDITS COMPLETE")
print(SEP)
print(f"Original : {SRC}")
print(f"Edited   : {DST}")
print()
print(f"Changes applied ({sum(1 for s, _ in log if s == 'DONE')}):")
for status, label in log:
    marker = "[DONE]" if status == "DONE" else "[MISS]"
    print(f"  {marker} {label}")

print()
print("=" * 65)
print("FLAGS -- AUTHOR ACTION REQUIRED (not auto-corrected)")
print("=" * 65)

print("""
1. TABLE S1 -- RURAL RESIDENCE (MEN): PROBABLE COPY-PASTE ERROR
   Row 'Rural residence', Men columns shows identical values to
   'Currently married' row:
     Low HGS:   81.1% [77.3%, 84.8%]
     High HGS:  90.6% [89.0%, 92.1%]
     Diff:      9.5%***
   These figures are implausible for rural residence (main Table 1
   shows overall male rural residence ~64%). Please replace with the
   correct values from your analysis output.

2. REFERENCE NUMBERING -- SUPPLEMENTARY vs MAIN TEXT MISMATCH
   The supplementary uses different reference numbers than the main
   manuscript. If both share a single reference list, the following
   need correction:

   Para 5 (GBTM recommendations): cited as [30]
     Main text uses [25] for Nagin & Odgers (2010) -- the GBTM
     recommendations paper. Reference [30] in the main text is a
     neuroimaging paper (Lee et al., unrelated to GBTM).
     --> Verify and likely change [30] to [25]

   Para 7 (kml package): cited as [26]
     Main text uses [27] for Genolini et al. (2015) -- kml package.
     --> Verify and likely change [26] to [27]

   Para 8 (clusterMLD package): cited as [27]
     Main text uses [28] for Zhou et al. (2023) -- clusterMLD.
     --> Verify and likely change [27] to [28]
""")

print("NEXT STEP -- view tracked changes in Word:")
print("  Review > Compare > Compare Documents")
print("  Original : SupplementaryData_AgeAgeing_v0418.docx")
print("  Revised  : SupplementaryData_AgeAgeing_v0418_edited.docx")
