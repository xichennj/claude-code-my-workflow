"""
Generate a referee report for the GripCog manuscript submitted to Age and Ageing.
Exports to a formatted Word document.
"""

import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\xc77\Dropbox\Grip Strength and Cognitive Trajectories\ReviewReport_AgeAgeing_GripCog.docx"

doc = docx.Document()

# -- Page margins -------------------------------------------------------------
for section in doc.sections:
    section.top_margin    = Inches(1.1)
    section.bottom_margin = Inches(1.1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

BODY_FONT = "Times New Roman"

# -- Helpers ------------------------------------------------------------------

def set_font(run, name=BODY_FONT, size=11, bold=False, italic=False, color=None):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    set_font(run, size=12, bold=True)
    return p


def body(text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_font(run)
    return p


def numbered(n, label_bold, text_rest):
    """Numbered comment with optional bold preamble then regular text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after       = Pt(8)
    p.paragraph_format.space_before      = Pt(3)
    p.paragraph_format.left_indent       = Inches(0.40)
    p.paragraph_format.first_line_indent = Inches(-0.40)
    r1 = p.add_run(f"{n}.  ")
    set_font(r1, bold=True)
    if label_bold:
        r2 = p.add_run(label_bold + "  ")
        set_font(r2, bold=True)
    r3 = p.add_run(text_rest)
    set_font(r3)
    return p


def bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Inches(0.40)
    r1 = p.add_run("•  ")
    set_font(r1, bold=True)
    r2 = p.add_run(text)
    set_font(r2)
    return p


def rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "888888")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


# =============================================================================
# TITLE BLOCK
# =============================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
set_font(p.add_run("REFEREE REPORT"), size=13, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
set_font(p.add_run("Age and Ageing"), size=11, italic=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
set_font(p.add_run(
    "“Sex-Specific Handgrip Strength Trajectories and Cognitive Performance "
    "in Older Chinese Adults: A Longitudinal Cohort Study”"
), size=11)

rule()

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(6)
r1 = p.add_run("Recommendation:  ")
set_font(r1, bold=True)
r2 = p.add_run("Major Revision")
set_font(r2, bold=True, color=(160, 0, 0))

rule()

# =============================================================================
# SUMMARY
# =============================================================================
h1("Summary")

body(
    "This paper uses group-based trajectory modelling (GBTM) on three waves of handgrip "
    "strength (HGS) data from CHARLS (2011–2015) to classify 8,714 middle-aged and older "
    "Chinese adults into High and Low HGS trajectory groups, then estimates the association of "
    "group membership with cognitive performance assessed in 2018. Analyses are conducted "
    "separately by sex, and trajectory classification is validated against two non-parametric "
    "clustering algorithms. The authors report that higher HGS trajectories are associated with "
    "better global cognition, episodic memory, and orientation in women, and with better "
    "visuospatial performance in men, with more pronounced effects in adults aged ≥60 years."
)
body(
    "The research question is timely and the dataset is well suited to address it. The "
    "sex-stratified, multi-method design represents a meaningful advance over prior cross-sectional "
    "or single-wave studies. However, several methodological concerns limit the strength of the "
    "conclusions that can currently be drawn. Two issues in particular—the absence of a "
    "formal sex-moderation test and the failure to establish that trajectory information adds "
    "value beyond baseline HGS level—are important enough to require new analyses before "
    "the paper can be recommended for publication."
)

# =============================================================================
# MAJOR COMMENTS
# =============================================================================
h1("Major Comments")

numbered(1,
    "Trajectory information versus baseline level.",
    "The paper’s central motivation is that longitudinal HGS trajectories capture "
    "information that cross-sectional or single-measurement approaches cannot. Yet Figure 2 "
    "shows that both the High and Low groups decline roughly in parallel across the three waves, "
    "with the primary distinction being initial level rather than rate of change. If trajectory "
    "group membership is essentially a restatement of average baseline HGS, the added value of "
    "repeated measurement is unclear. The authors should formally test whether HGS trajectory "
    "group membership predicts Wave 4 cognition beyond what a single baseline HGS value "
    "predicts—for example, by adding Model 3 specifications in which observed baseline HGS "
    "(continuous) replaces trajectory group membership, and comparing standardised effect "
    "sizes. Without this test, the rationale for the trajectory-modelling framework, as opposed "
    "to a simpler baseline-measurement approach, is not established."
)

numbered(2,
    "Absence of a formal sex-moderation test.",
    "The paper compares sex-stratified results and concludes that associations are "
    "‘sex-specific’ and reflect ‘distinct sex-specific cognitive phenotypes.’ "
    "However, this conclusion rests on an informal comparison of two separate models, not a "
    "formal interaction test. A pooled model including a sex × HGS-group interaction term "
    "is needed to determine whether the between-sex differences in effect pattern are "
    "statistically distinguishable from chance variation across subgroups. This is particularly "
    "important given that the visuospatial finding in men and the orientation finding in women "
    "both have confidence intervals that include zero or near-zero at conventional thresholds. "
    "The interaction analysis should be added, at minimum as a supplementary table."
)

numbered(3,
    "Multiple comparisons and the prominence of the exploratory caveat.",
    "The paper tests five cognitive outcomes in two sex strata, yielding ten primary comparisons "
    "per model. Under a Bonferroni-corrected threshold of P < 0.005, several domain-specific "
    "associations—notably orientation in women (P ≈0.045) and visuospatial function in "
    "women (P ≈0.053, not significant)—would not survive correction. The paper "
    "acknowledges exploratory status in a single sentence, but this caveat must be more "
    "prominent. Specifically: (a) the Abstract should not present the orientation finding as a "
    "confirmed result without qualification; (b) the Key Points bullets should be softened "
    "accordingly; and (c) the Discussion should state explicitly that domain-specific findings "
    "require independent replication before conclusions about distinct cognitive phenotypes "
    "can be drawn."
)

numbered(4,
    "Justification and visualisation of the two-group solution.",
    "The BIC-based selection of a two-group model is described narratively but the BIC values "
    "are never shown. A figure plotting BIC (or ΔBIC) against the number of groups (one to "
    "six) is standard practice in GBTM studies and is required by the GRoLTS reporting "
    "guidelines (reference 18 in the manuscript). This figure should be added to the "
    "Supplementary Data. Additionally, for a sample of approximately 4,760 women and 3,954 men, "
    "a two-group solution may be overly parsimonious. The authors should report whether "
    "three- or four-group models revealed clinically distinct subgroups (e.g., a ‘rapidly "
    "declining’ group) that were ultimately collapsed, and provide fuller justification "
    "for retaining only two groups beyond the BIC elbow."
)

numbered(5,
    "Missing data and the plausibility of the MAR assumption.",
    "Approximately 27% of eligible respondents were excluded (11,988 → 8,714), "
    "predominantly for missing HGS or cognitive data. Inverse-probability weighting under "
    "missing at random (MAR) is applied, but no attrition-bias diagnostics are presented. "
    "The following should be added: (a) a comparison of baseline characteristics between "
    "included and excluded participants; (b) the range and distribution of IPW weights, "
    "including whether extreme weights were trimmed or stabilised; and (c) an explicit "
    "acknowledgement that MAR is unlikely to hold if individuals with concurrent physical and "
    "cognitive impairment are preferentially lost—in which case IPW would underestimate "
    "the true association. If feasible, a sensitivity analysis using pattern mixture modelling "
    "or Rosenbaum-style bounds would strengthen the paper."
)

numbered(6,
    "Cognitive Z-score standardisation.",
    "The paper states that domain-specific Z-scores were the outcomes but does not specify "
    "whether standardisation used the mean and SD of the full analytic sample, within-sex "
    "distributions, or an external reference. Given the large sex gap in raw cognitive scores "
    "(global score 13.3 versus 15.5), a one-SD shift in a pooled distribution represents "
    "a different absolute change for women and men, complicating cross-sex comparisons of "
    "effect sizes. The standardisation approach must be stated explicitly, and if pooled "
    "Z-scores were used, the authors should confirm that this does not distort the sex "
    "comparisons."
)

numbered(7,
    "The 22 kg and 33 kg reference values and their clinical framing.",
    "A Key Points bullet states that ‘a model-estimated intercept of approximately 22 kg "
    "in women may serve as an empirical reference point for research benchmarking.’ The "
    "Asian Working Group for Sarcopenia (AWGS 2019) already defines possible sarcopenia "
    "thresholds at <20 kg (women) and <28 kg (men) for Asian populations. The Low-group mean "
    "of 22 kg in women sits between the AWGS possible and probable sarcopenia cut-offs, and "
    "the 33 kg mean in men falls above the AWGS criteria entirely. Presenting model-estimated "
    "group means as reference points for screening—even with caveats—without "
    "contextualising them against AWGS 2019 risks misrepresenting their clinical standing. "
    "The Key Points bullet should be removed or rewritten to make clear that these are "
    "descriptive group means in this particular sample; and the Discussion should include "
    "explicit comparison with AWGS 2019 thresholds."
)

# =============================================================================
# MINOR COMMENTS
# =============================================================================
h1("Minor Comments")

numbered(1,
    "Top-coding of HGS at 60 kg.",
    "The Methods state that HGS was top-coded at 60 kg, but Figure 2 shows group means "
    "well below 40 kg, suggesting this ceiling is rarely binding. The number of participants "
    "affected and the rationale for 60 kg as the ceiling should be stated."
)

numbered(2,
    "Pentagon copying labelled as visuospatial function.",
    "The visuospatial score derives from a single item—copying overlapping pentagons. "
    "This task primarily captures constructional ability and fine motor control, and "
    "its validity as a measure of visuospatial processing broadly defined is limited. "
    "The label ‘visuospatial function’ should be qualified throughout "
    "(e.g., ‘constructional visuospatial ability’), and the Discussion of "
    "the sex-specific visuospatial finding in men should acknowledge this constraint explicitly."
)

numbered(3,
    "Clinical significance of the orientation finding in women.",
    "The paper reports a significant HGS–orientation association in women (P ≈0.045) "
    "but does not contextualise it clinically. Orientation deficits are a hallmark of early "
    "dementia—and specifically of Alzheimer-type pathology—rather than of normal "
    "cognitive ageing. If this association is robust, it carries more serious clinical "
    "implications than an episodic memory association of similar magnitude, and the "
    "Discussion should address this distinction directly."
)

numbered(4,
    "Age mixing across midlife and older age.",
    "The analytic sample spans ages 45–80. HGS–cognition relationships differ "
    "substantially between midlife (45–64) and older adulthood (≥65). Although a "
    "≥60 subgroup analysis is provided, the main models pool these groups. At minimum, "
    "the authors should confirm that the two-group trajectory structure is stable within "
    "the 45–59 and ≥60 age bands, or discuss how age heterogeneity may affect "
    "the trajectory estimates."
)

numbered(5,
    "Kendall’s W interpretation.",
    "The Supplementary Data reports Kendall’s W of 0.80–0.81. While polychoric "
    "correlations above 0.85 are described as indicating ‘exceptional concordance,’ "
    "Kendall’s W of 0.80–0.81 indicates substantial but not exceptional agreement "
    "(values ≥0.90 are typically required for strong concordance). The language "
    "in the Supplementary Data should be calibrated to reflect this distinction."
)

numbered(6,
    "IPW construction details.",
    "The paper does not state whether the IPW weights were stabilised (using the marginal "
    "probability of inclusion in the numerator) or left as raw inverse probabilities, and "
    "whether any truncation was applied. Unstabilised weights can have high variance and "
    "unduly influential observations. These details should be reported in the Methods."
)

numbered(7,
    "Form of robust standard errors.",
    "The paper reports ‘robust standard errors’ without specifying the form. "
    "Given CHARLS’s complex multi-stage clustered sampling design, cluster-robust "
    "standard errors accounting for within-community correlation are likely more appropriate "
    "than heteroskedasticity-consistent (HC) standard errors. The authors should clarify "
    "which form was used and whether community-level clustering was accounted for."
)

numbered(8,
    "Reference Wu et al. 2026 (reference 29).",
    "This citation carries a 2026 journal year. If the article was published online ahead "
    "of the 2026 issue print date, the online-first publication date should be given. "
    "If the article is forthcoming, its use as a citation should be reconsidered."
)

# =============================================================================
# SPECIFIC / TEXTUAL COMMENTS
# =============================================================================
h1("Specific and Textual Comments")

numbered(1,
    "Abstract, Methods.",
    "The three specific measurement years (2011, 2013, 2015) should be stated rather than "
    "‘three waves from 2011–2015,’ as the inter-wave interval is relevant "
    "for interpreting trajectory slope estimates."
)

numbered(2,
    "Introduction, paragraph 3.",
    "The sentence ‘These gaps call for a longitudinal design that addresses sex heterogeneity "
    "and evaluates trajectory classification stability across multiple methods’ reads as a "
    "description of the present study’s design rather than a gap in the literature. "
    "Revise to separate the identified gap from the proposed solution."
)

numbered(3,
    "Table 1 footnote.",
    "The note should state clearly that all Diff. values represent Women minus Men. "
    "The current wording (‘Women minus Men for categorical, Men minus Women for continuous "
    "variables’) is inconsistent with the actual values in the table."
    " [Corrected in the edited manuscript file; please verify in the final submission.]"
)

numbered(4,
    "Supplementary Table S1, Rural residence (Men).",
    "The Men columns for Rural residence show values identical to the Currently married row "
    "(81.1% / 90.6% / 9.5%***), which is inconsistent with the overall male rural residence "
    "of ~64% in main Table 1. This is a probable copy-paste error and must be corrected "
    "with values from the original analysis output before submission."
)

numbered(5,
    "Supplementary Methods, reference numbers.",
    "Citations [26], [27], and [30] in the Supplementary Methods appear to use different "
    "numbers than the corresponding entries in the main text reference list "
    "(kml package = [27]; clusterMLD = [28]; Nagin & Odgers = [25]). "
    "Please verify and harmonise across the full document."
)

numbered(6,
    "Discussion, effect size paragraph.",
    "Two successive sentences both use the construction ‘should be read as’, "
    "creating awkward repetition. One should be rephrased."
)

numbered(7,
    "Discussion, paragraph 3.",
    "‘evidence from China similarly suggest’ should be ‘suggests’ "
    "(subject–verb agreement)."
)

# =============================================================================
# SUMMARY OF REQUIRED ACTIONS
# =============================================================================
rule()
h1("Summary of Required Actions for Revision")

body("The following items must be addressed before the manuscript can be reconsidered:")

required = [
    "New analysis: formal sex x HGS-group interaction test in a pooled model (Major Comment 2).",
    "New analysis: comparison of trajectory-group vs baseline HGS as a continuous predictor (Major Comment 1).",
    "New supplementary figure: BIC values across one- to six-group solutions (Major Comment 4).",
    "New supplementary table or text: attrition-bias diagnostics and IPW weight distribution (Major Comment 5).",
    "Revise Abstract and Key Points to increase prominence of multiple-comparisons caveat (Major Comment 3).",
    "Remove or substantially revise the 22 kg / 33 kg Key Points bullet; add AWGS 2019 comparison in Discussion (Major Comment 7).",
    "State cognitive Z-score standardisation approach explicitly (Major Comment 6).",
    "Correct Rural residence copy-paste error in Supplementary Table S1 (Specific Comment 4).",
    "Verify and harmonise reference numbers between main text and Supplementary Data (Specific Comment 5).",
    "Clarify IPW stabilisation and weight truncation in Methods (Minor Comment 6).",
    "Clarify form of robust standard errors; confirm CHARLS cluster structure accounted for (Minor Comment 7).",
]

for item in required:
    bullet(item)

doc.save(OUT)
print(f"Saved: {OUT}")
