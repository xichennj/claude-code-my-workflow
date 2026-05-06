"""
Polish GripCog_AgeAgeing_v0418.docx for Age and Ageing submission.
Saves edited copy; use Word Review > Compare Documents to view tracked changes.
"""

import shutil
import docx

SRC = r"C:\Users\xc77\Dropbox\Grip Strength and Cognitive Trajectories\GripCog_AgeAgeing_v0418.docx"
DST = r"C:\Users\xc77\Dropbox\Grip Strength and Cognitive Trajectories\GripCog_AgeAgeing_v0418_edited.docx"

shutil.copy(SRC, DST)
doc = docx.Document(DST)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def full_text(para):
    return "".join(r.text for r in para.runs)


def replace_in_para(para, old, new):
    """Replace old->new in para, consolidating all runs into the first.
    Returns True if replacement was made."""
    ft = full_text(para)
    if old not in ft:
        return False
    new_ft = ft.replace(old, new, 1)
    if para.runs:
        para.runs[0].text = new_ft
        for r in para.runs[1:]:
            r.text = ""
    return True


log = []

def edit(para, old, new, label):
    if replace_in_para(para, old, new):
        log.append(("DONE", label))
        return True
    return False


# ---------------------------------------------------------------------------
# Collect all paragraphs (body + inside tables)
# ---------------------------------------------------------------------------
all_paras = list(doc.paragraphs)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            all_paras.extend(cell.paragraphs)

# ---------------------------------------------------------------------------
# Apply edits
# ---------------------------------------------------------------------------
for para in all_paras:
    ft = full_text(para)

    # ---- ABSTRACT --------------------------------------------------------
    # More precise wave specification (matches pandoc output)
    edit(para,
         "HGS was measured across three waves from 2011-2015",
         "HGS was measured at three time points (2011, 2013, and 2015)",
         "Abstract: clarified wave specification")

    # ---- INTRODUCTION ----------------------------------------------------
    # Double space before "Third" in knowledge-gaps paragraph
    edit(para,
         "within-sex heterogeneity [17].  Third",
         "within-sex heterogeneity [17]. Third",
         "Intro: removed extra space before 'Third'")

    # Spurious comma before coordinated verb clause
    edit(para,
         "sex heterogeneity, and evaluates trajectory classification stability",
         "sex heterogeneity and evaluates trajectory classification stability",
         "Intro: removed spurious comma before 'and evaluates'")

    # ---- METHODS / DATA --------------------------------------------------
    # Grammar + spacing fix: double space + incorrect verb form
    edit(para,
         "only have  one HGS measure",
         "with only one HGS measurement",
         "Methods/Data: grammar fix ('only have  one HGS measure' -> 'with only one HGS measurement')")

    # British English spelling
    edit(para,
         "CHARLS enrolls adults from age 45",
         "CHARLS enrols adults from age 45",
         "Methods/Data: British English 'enrolls' -> 'enrols'")

    # ---- METHODS / MODELS ------------------------------------------------
    # Sentence fragment "solution. and" + hyphenation
    edit(para,
         "We tested one to six groups solution. and the optimal number of groups was determined by identifying the elbow point using the Bayesian Information Criterion (BIC) curve.",
         "We tested one- to six-group solutions, and the optimal number of groups was selected at the elbow of the Bayesian Information Criterion (BIC) curve.",
         "Methods/Models: fixed sentence fragment, hyphenation, BIC wording")

    # Missing space after period
    edit(para,
         "compared.Quadratic",
         "compared. Quadratic",
         "Methods/Models: added missing space after period before 'Quadratic'")

    # Informal sentence-initial "Then"
    edit(para,
         "Then, weighted linear regression examined associations",
         "Weighted linear regression then examined associations",
         "Methods/Models: moved 'then' out of sentence-initial position")

    # Preposition: "on" -> "based on"
    edit(para,
         "predicting inclusion on baseline age",
         "predicting inclusion based on baseline age",
         "Methods/Models: preposition fix 'on' -> 'based on'")

    # Preposition + spacing: "with" + "CHARLSsurvey"
    edit(para,
         "multiplied it with the baseline CHARLSsurvey weight",
         "multiplied it by the baseline CHARLS survey weight",
         "Methods/Models: 'with' -> 'by'; added space in 'CHARLS survey'")

    # ---- RESULTS ---------------------------------------------------------
    # Duplicate "substantial"
    edit(para,
         "attenuated substantially, indicating substantial confounding",
         "attenuated substantially, indicating considerable confounding",
         "Results: removed duplicate 'substantial'")

    # ---- DISCUSSION ------------------------------------------------------
    # Repetitive "should be read as ... should be read as"
    # Uses actual Unicode em dash (U+2014)
    edit(para,
         ("so this magnitude should be read as a population-scale signal rather than an individual-level effect. "
          "Our estimates fall within the range reported by coordinated multi-study analyses of HGS and cognition in older adults [8]. "
          "Rather than marking a threshold for intervention at the individual level, these magnitudes should be read as signalling "
          "population-level variation in cognitive vulnerability that accumulates where low HGS is prevalent."),
         ("and our estimates fall within the range reported by coordinated multi-study analyses of HGS and cognition "
          "in older adults [8]. These magnitudes are best interpreted as population-level signals of cognitive "
          "vulnerability rather than thresholds for individual-level intervention."),
         "Discussion: consolidated repetitive 'should be read as' phrasing")

    # Subject-verb agreement
    edit(para,
         "evidence from China similarly suggest that",
         "evidence from China similarly suggests that",
         "Discussion: subject-verb agreement 'suggest' -> 'suggests'")

    # ---- SECTION HEADING -------------------------------------------------
    if ft.strip() == "Acknowledgement":
        if para.runs:
            para.runs[0].text = ft.replace("Acknowledgement", "Acknowledgements", 1)
            for r in para.runs[1:]:
                r.text = ""
        log.append(("DONE", "Heading: 'Acknowledgement' -> 'Acknowledgements'"))

    # ---- TABLE 1 FOOTNOTE ------------------------------------------------
    edit(para,
         "Women minus Men for categorical variables, Men minus Women for continuous variables",
         "Women minus Men",
         "Table 1 note: corrected Diff. formula (all values are Women minus Men)")


# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(DST)

# ---------------------------------------------------------------------------
# Report
# ---------------------------------------------------------------------------
SEP = "-" * 65
print(SEP)
print("EDITING COMPLETE")
print(SEP)
print(f"Original: {SRC}")
print(f"Edited  : {DST}")
print()
print(f"Changes applied ({sum(1 for s,_ in log if s=='DONE')}):")
for status, label in log:
    print(f"  [{status}] {label}")

print()
print("FLAG - AUTHOR ACTION REQUIRED")
print("  Numerical discrepancy between text and Table 2 for")
print("  Men, global cognition, Model 2:")
print("    Body text : beta = 0.042; 95% CI: -0.020 to 0.104")
print("    Table 2   : beta = 0.043; 95% CI: -0.019 to 0.105")
print("  Please verify against analysis output and make consistent.")

print()
print("NEXT STEP - view tracked changes in Word:")
print("  Review tab > Compare > Compare Documents")
print("  Original document : GripCog_AgeAgeing_v0418.docx")
print("  Revised document  : GripCog_AgeAgeing_v0418_edited.docx")
