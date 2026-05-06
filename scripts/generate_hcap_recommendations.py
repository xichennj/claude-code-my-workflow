"""
Export major revision recommendations for the HCAP dementia abstract to Word.
"""

import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\xc77\Dropbox\Algorithmic Performance and Fairness of Dementia in HCAP_recommendations.docx"

doc = docx.Document()

for section in doc.sections:
    section.top_margin    = Inches(1.1)
    section.bottom_margin = Inches(1.1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

FONT = "Times New Roman"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name   = FONT
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(text, size=13, space_before=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text.upper())
    set_font(r, size=size, bold=True)
    return p


def subheading(number, severity_label, severity_color, title):
    """e.g.  '1.  REFRAME THE STUDY TYPE  [Critical]' """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(f"{number}.  {title}  ")
    set_font(r1, size=11, bold=True)
    r2 = p.add_run(f"[{severity_label}]")
    set_font(r2, size=11, bold=True, color=severity_color)
    return p


def body(text, space_after=6, left_indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if left_indent:
        p.paragraph_format.left_indent = Inches(left_indent)
    r = p.add_run(text)
    set_font(r)
    return p


def label_para(label, text):
    """Bold label followed by regular text on the same line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.25)
    r1 = p.add_run(label + "  ")
    set_font(r1, bold=True)
    r2 = p.add_run(text)
    set_font(r2, italic=True)
    return p


def rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "888888")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


# Severity colours
RED    = (180, 0,   0)
ORANGE = (180, 90,  0)
GRAY   = (100, 100, 100)

# ---------------------------------------------------------------------------
# Title block
# ---------------------------------------------------------------------------
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
set_font(p.add_run("ABSTRACT REVIEW: MAJOR RECOMMENDATIONS"), size=13, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
set_font(p.add_run("JAMA Submission"), size=11, italic=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
set_font(p.add_run(
    "“Algorithmic Performance and Fairness of Machine Learning Models for "
    "Dementia Classification Using the Harmonized Cognitive Assessment Protocol (HCAP)”"
), size=11)

rule()

# ---------------------------------------------------------------------------
# Overview
# ---------------------------------------------------------------------------
heading("Overview", space_before=8)

body(
    "The abstract reports a well-motivated study with a strong dataset and a relevant "
    "public-health question. However, several framing and methodological gaps would likely "
    "trigger desk rejection or an immediate major revision request at JAMA. The three critical "
    "items below (study-type label, reference standard, and train/test methodology) should be "
    "addressed before submission. The remaining items are important for passing peer review."
)

# ---------------------------------------------------------------------------
# Recommendations
# ---------------------------------------------------------------------------
heading("Recommendations")

# ── 1 ───────────────────────────────────────────────────────────────────────
subheading(1, "Critical", RED, "Reframe the Study Type")

body(
    "The Design sentence calls this a “cross-sectional study.” Cross-sectional describes "
    "the observational structure of the underlying data, not the analytical task. What the paper "
    "actually performs is prediction model development and internal validation — a distinct "
    "study type that JAMA now routes through TRIPOD-AI reporting requirements. Labelling it "
    "cross-sectional signals to editors and reviewers that the paper may not have followed the "
    "appropriate reporting framework."
)
label_para("Where to fix:", "Design, Setting, and Participants")
label_para("Recommended change:",
    "“This cross-sectional study used…”  →  "
    "“This prediction model development and validation study used…”  "
    "Confirm the paper follows the TRIPOD-AI checklist.")

# ── 2 ───────────────────────────────────────────────────────────────────────
subheading(2, "Critical", RED, "Name the Reference Standard for Dementia")

body(
    "The abstract trains models on “clinician-validated dementia classifications” but "
    "never states what that standard is. In classification studies the reference standard "
    "is as important as the outcome itself — its sensitivity and specificity set the ceiling "
    "for any downstream algorithm. JAMA reviewers will flag this immediately. Is it the "
    "Langa-Weir classification, an expert consensus diagnosis, or a formal clinical workup?"
)
label_para("Where to fix:", "Design, Setting, and Participants")
label_para("Recommended change:",
    "Name the standard explicitly, e.g., “…using the Langa-Weir clinician-adjudicated "
    "dementia classification as the reference standard.”")

# ── 3 ───────────────────────────────────────────────────────────────────────
subheading(3, "Critical", RED, "State the Train/Test Methodology")

body(
    "The abstract reports accuracy, sensitivity, and specificity figures but says nothing about "
    "how the models were trained and tested. Were these from 5-fold cross-validation, a held-out "
    "test set, or the full sample? Without this, the metrics are uninterpretable: internal "
    "cross-validation estimates are optimistic, and full-sample estimates are uninformative "
    "about generalisability. This is a mandatory element for any prediction model paper."
)
label_para("Where to fix:", "Main Outcomes and Measures or Design")
label_para("Recommended change:",
    "Add one sentence, e.g., “Models were trained using 5-fold cross-validation on a "
    "randomly selected 80% training set and evaluated on a held-out 20% test set” "
    "(adjust to match the actual approach).")

# ── 4 ───────────────────────────────────────────────────────────────────────
subheading(4, "Important", ORANGE, "Revise the Objective to Include “Development”")

body(
    "The Objective reads “To evaluate predictive performance and algorithmic fairness…” "
    "Evaluate implies auditing pre-existing models. If the paper trains new models, the "
    "Objective should say so. This also affects how reviewers assess novelty: developing new "
    "algorithms is a distinct contribution from evaluating existing ones."
)
label_para("Where to fix:", "Objective")
label_para("Recommended change:",
    "“To evaluate…”  →  “To develop and evaluate…”")

# ── 5 ───────────────────────────────────────────────────────────────────────
subheading(5, "Important", ORANGE,
           "Quantify the Comparison with Existing Algorithms")

body(
    "The Results state “these models demonstrated substantially improved case detection” "
    "compared with existing HRS algorithms, but provide no figures. JAMA does not accept "
    "unquantified comparative claims in Results. Readers need to know the baseline sensitivity "
    "of the existing algorithm and the magnitude of the improvement to judge whether it "
    "is clinically meaningful."
)
label_para("Where to fix:", "Results")
label_para("Recommended change:",
    "Replace the vague comparison with a specific figure, e.g., “Sensitivity improved "
    "from X% with the existing Langa-Weir algorithm to a range of 0.81–0.93 across "
    "the new models.”")

# ── 6 ───────────────────────────────────────────────────────────────────────
subheading(6, "Important", ORANGE,
           "Clarify the Direction of Fairness Violations")

body(
    "The Key Points state Hispanic and male participants showed lower “equal opportunity "
    "and predictive equality.” These are opposite types of error: equal opportunity is the "
    "true positive rate (missed dementia cases), predictive equality is the false positive rate "
    "(false alarms). If a group shows lower rates on both simultaneously, that indicates "
    "systematic under-detection across the board — a distinct and more serious pattern than "
    "a single-metric violation. The abstract conflates these without clarification."
    "\n\n"
    "Additionally, the Conclusions say disparities appear “in several models” but do not "
    "indicate whether any model achieves both high performance and acceptable fairness. "
    "This tradeoff is typically the central actionable finding in algorithmic fairness work."
)
label_para("Where to fix:", "Key Points – Findings; Results; Conclusions")
label_para("Recommended change:",
    "Specify which fairness metrics are violated for which groups, and in which direction. "
    "In the Conclusions, indicate whether any model approximates the fairness-performance "
    "frontier, or explicitly note that no tested model satisfied both criteria.")

# ── 7 ───────────────────────────────────────────────────────────────────────
subheading(7, "Important", ORANGE,
           "Acknowledge the Absence of External Validation")

body(
    "All reported performance metrics appear to derive from a single dataset (HRS-HCAP 2016). "
    "Even rigorous cross-validation constitutes internal validation only — performance in "
    "a new cohort, a different time period, or a different healthcare system is unknown. "
    "JAMA reviewers evaluating a prediction model paper will expect this limitation to be "
    "acknowledged explicitly in the abstract."
)
label_para("Where to fix:", "Conclusions and Relevance")
label_para("Recommended change:",
    "Add to the final sentence: “…findings require external validation in independent "
    "cohorts before clinical or research deployment.”")

# ── 8 ───────────────────────────────────────────────────────────────────────
subheading(8, "Minor", GRAY,
           "Clarify Super Learner’s Relationship to the Other Five Models")

body(
    "Super Learner is an ensemble method that typically uses other algorithms as base learners. "
    "If the five other models (logistic regression, LASSO, RF, SVM, XGBoost) serve as its "
    "components, framing the comparison as “six independent algorithms” is potentially "
    "misleading, and a reviewer familiar with ensemble methods will raise this. The relationship "
    "should be clarified in one sentence."
)
label_para("Where to fix:", "Main Outcomes and Measures")
label_para("Recommended change:",
    "Add a parenthetical clarification, e.g., “…and Super Learner (an ensemble of "
    "the preceding five algorithms).”")

# ---------------------------------------------------------------------------
# Summary table
# ---------------------------------------------------------------------------
rule()
heading("Summary", space_before=8)

body("Items 1–3 are most likely to trigger desk rejection and should be resolved before submission.")

table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"

hdr = table.rows[0].cells
for cell, txt in zip(hdr, ["#", "Recommendation", "Severity", "Section"]):
    p = cell.paragraphs[0]
    r = p.add_run(txt)
    set_font(r, bold=True, size=10)
    p.paragraph_format.space_after = Pt(2)

rows = [
    ("1", "Reframe as prediction model development study",   "Critical",  "Design"),
    ("2", "Name the reference standard for dementia",        "Critical",  "Design"),
    ("3", "State train/test methodology",                    "Critical",  "Main Outcomes"),
    ("4", "Add ‘development’ to the Objective",    "Important", "Objective"),
    ("5", "Quantify comparison with existing algorithms",    "Important", "Results"),
    ("6", "Clarify direction of fairness violations",        "Important", "KP + Results + Conclusions"),
    ("7", "Acknowledge absence of external validation",      "Important", "Conclusions"),
    ("8", "Clarify Super Learner vs constituent models",     "Minor",     "Main Outcomes"),
]

severity_colors = {
    "Critical":  RGBColor(180, 0,  0),
    "Important": RGBColor(180, 90, 0),
    "Minor":     RGBColor(100, 100, 100),
}

for num, rec, sev, section in rows:
    row = table.add_row().cells
    for cell, txt in zip(row, [num, rec, sev, section]):
        p = cell.paragraphs[0]
        r = p.add_run(txt)
        color = severity_colors.get(sev) if txt == sev else None
        set_font(r, size=10, bold=(txt == sev), color=color)
        p.paragraph_format.space_after = Pt(2)

# Column widths
widths = [Inches(0.3), Inches(3.0), Inches(1.0), Inches(1.5)]
for row in table.rows:
    for cell, w in zip(row.cells, widths):
        cell.width = w

doc.save(OUT)
print(f"Saved: {OUT}")
